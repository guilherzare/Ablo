"""
Export du bilan en Word (.docx) et PDF.
Émet des lignes JSON sur stdout pour le streaming de progression.
"""
import json
import datetime
import uuid
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor, black, white, lightgrey
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False

from settings_manager import get_settings
from patient_manager import patient_dir_for, record_bilan

AUTEVAL_CRITERIA = [
    "État initial",
    "Envie de revenir",
    "Bien fait",
    "Beau",
    "Bon moment",
    "État final",
]

DOT_LABELS = ["—", "Faible", "Passable", "Bien", "Très bien", "Excellent"]


def _parse_auteval(content: str) -> dict | None:
    """Retourne None si ce n'est pas de l'autoéval, sinon un dict typé."""
    try:
        parsed = json.loads(content)
        if isinstance(parsed, dict):
            if parsed.get("type") == "multi_session":
                return parsed
            else:
                return {"type": "single", "scores": parsed}
    except Exception:
        pass
    return None


def _emit(payload: dict) -> None:
    print(json.dumps(payload, ensure_ascii=False), flush=True)


def _output_dir() -> Path:
    settings = get_settings()
    folder = settings.get("export_folder", "~/Documents/Ablo")
    path = Path(folder).expanduser()
    path.mkdir(parents=True, exist_ok=True)
    return path


def _make_filename(session_id: str) -> str:
    return f"bilan_{datetime.date.today().strftime('%Y%m%d')}_{session_id}"


def _format_date(iso: str) -> str:
    try:
        return datetime.datetime.fromisoformat(iso).strftime("%d/%m/%Y")
    except Exception:
        return iso[:10] if len(iso) >= 10 else iso


def _export_docx(
    sections: list[dict],
    template_name: str,
    patient_name: str,
    out_dir: Path,
    filename: str,
    settings: dict,
    patient_label: str = "",
) -> Path:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    therapist_name = settings.get("therapist_name", "")
    therapist_email = settings.get("therapist_email", "")
    date_str = datetime.date.today().strftime("%d/%m/%Y")

    # En-tête : nom et email thérapeute sur la première page uniquement
    section0 = doc.sections[0]
    section0.different_first_page_header_footer = True
    hp = section0.first_page_header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if therapist_name:
        run = hp.add_run(therapist_name)
        run.bold = True
        run.font.size = Pt(10)
    if therapist_email:
        hp.add_run(f"\n{therapist_email}").font.size = Pt(9)

    # Titre principal
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(12)
    title_run = title_para.add_run("Bilan séances Art-thérapie")
    title_run.bold = True
    title_run.font.size = Pt(14)

    # Nom du patient — même style que le titre
    if patient_name:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(patient_name)
        sub_run.bold = True
        sub_run.font.size = Pt(14)

    doc.add_paragraph()

    # Sections
    for s in sections:
        heading = doc.add_heading(s["title"], level=1)
        heading.runs[0].font.size = Pt(11)
        heading.runs[0].bold = True

        content = s.get("content", "").strip()
        auteval = _parse_auteval(content) if "autoévaluation" in s["title"].lower() else None

        if auteval is not None:
            if auteval["type"] == "multi_session":
                sessions = auteval["sessions"]
                header_row = ["Critère"] + [
                    f"S{i+1}\n{_format_date(sess.get('date', ''))}"
                    for i, sess in enumerate(sessions)
                ]
                table = doc.add_table(rows=1, cols=len(header_row))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for cell, text in zip(hdr_cells, header_row):
                    run = cell.paragraphs[0].add_run(text)
                    run.bold = True
                    run.font.size = Pt(9)
                for criterion in AUTEVAL_CRITERIA:
                    row_cells = table.add_row().cells
                    row_cells[0].paragraphs[0].add_run(criterion).font.size = Pt(9)
                    for i, sess in enumerate(sessions):
                        scores = sess.get("scores", {})
                        cell_text = "—" if not scores else str(scores.get(criterion, 0))
                        row_cells[i + 1].paragraphs[0].add_run(cell_text).font.size = Pt(9)
                legend = doc.add_paragraph()
                legend_run = legend.add_run(
                    "Scores de 0 à 5 — 0 : non renseigné · 1 : Faible · 2 : Passable · "
                    "3 : Bien · 4 : Très bien · 5 : Excellent. "
                    "La première séance ne comporte pas d'autoévaluation (—)."
                )
                legend_run.italic = True
                legend_run.font.size = Pt(8)
                legend_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
            else:
                scores = auteval["scores"]
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                for cell, text in zip(hdr, ["Critère", "Note /5", "Appréciation"]):
                    run = cell.paragraphs[0].add_run(text)
                    run.bold = True
                    run.font.size = Pt(9)
                for criterion in AUTEVAL_CRITERIA:
                    val = scores.get(criterion, 0)
                    row = table.add_row().cells
                    row[0].paragraphs[0].add_run(criterion).font.size = Pt(9)
                    row[1].paragraphs[0].add_run(str(val)).font.size = Pt(9)
                    label = DOT_LABELS[val] if 0 <= val < len(DOT_LABELS) else ""
                    row[2].paragraphs[0].add_run(label).font.size = Pt(9)

        elif content:
            for line in content.split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith(("- ", "* ", "• ")):
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(stripped[2:])
                else:
                    doc.add_paragraph(stripped)
        else:
            p = doc.add_paragraph()
            run = p.add_run("[Section non renseignée]")
            run.italic = True
            run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        doc.add_paragraph()

    # Formule de clôture — lieu du patient (tag) ou placeholder rouge si absent
    closing_para = doc.add_paragraph()
    closing_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    gray = RGBColor(0x37, 0x41, 0x51)

    r1 = closing_para.add_run("Merci de votre confiance, fait à ")
    r1.font.size = Pt(9)
    r1.font.color.rgb = gray

    if patient_label:
        r2 = closing_para.add_run(patient_label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = gray
    else:
        r2 = closing_para.add_run("[Lieu]")
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    r3 = closing_para.add_run(f", le {date_str}.")
    r3.font.size = Pt(9)
    r3.font.color.rgb = gray

    if therapist_name:
        r4 = closing_para.add_run(f"\n{therapist_name}")
        r4.font.size = Pt(9)
        r4.font.color.rgb = gray

    dest = out_dir / f"{filename}.docx"
    doc.save(str(dest))
    return dest


def _export_pdf(
    sections: list[dict],
    template_name: str,
    patient_name: str,
    out_dir: Path,
    filename: str,
    settings: dict,
    patient_label: str = "",
) -> Path:
    dest = out_dir / f"{filename}.pdf"

    therapist_name = settings.get("therapist_name", "")
    therapist_email = settings.get("therapist_email", "")
    date_str = datetime.date.today().strftime("%d/%m/%Y")

    styles = getSampleStyleSheet()
    indigo = HexColor("#4F46E5")

    title_style = ParagraphStyle(
        "BoldTitle", parent=styles["Title"],
        fontSize=14, spaceAfter=4,
    )
    # Nom du patient — même taille et couleur que le titre
    patient_style = ParagraphStyle(
        "PatientName", parent=styles["Normal"],
        fontSize=14, fontName="Helvetica-Bold", spaceAfter=16, alignment=1,
    )
    h1_style = ParagraphStyle(
        "SectionH1", parent=styles["Heading1"],
        fontSize=11, spaceBefore=14, spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=10, leading=15, spaceAfter=4,
    )
    small_style = ParagraphStyle(
        "Small", parent=styles["Normal"],
        fontSize=8, textColor=HexColor("#9CA3AF"), spaceAfter=2,
    )
    closing_style = ParagraphStyle(
        "Closing", parent=styles["Normal"],
        fontSize=9, textColor=HexColor("#374151"), alignment=2, spaceBefore=24,
    )

    def on_first_page(canvas, doc):
        canvas.saveState()
        if therapist_name:
            canvas.setFont("Helvetica-Bold", 9)
            canvas.drawString(2.5 * cm, A4[1] - 1.5 * cm, therapist_name)
        if therapist_email:
            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(HexColor("#6B7280"))
            canvas.drawString(2.5 * cm, A4[1] - 1.5 * cm - 12, therapist_email)
        canvas.restoreState()

    def on_later_pages(canvas, doc):
        pass

    doc = SimpleDocTemplate(
        str(dest), pagesize=A4,
        rightMargin=2.5 * cm, leftMargin=2.5 * cm,
        topMargin=3 * cm, bottomMargin=3 * cm,
    )

    story = []
    story.append(Paragraph("Bilan séances Art-thérapie", title_style))
    if patient_name:
        story.append(Paragraph(patient_name, patient_style))
    story.append(Spacer(1, 0.4 * cm))

    for s in sections:
        story.append(Paragraph(s["title"], h1_style))
        content = s.get("content", "").strip()
        auteval = _parse_auteval(content) if "autoévaluation" in s["title"].lower() else None

        if auteval is not None:
            if auteval["type"] == "multi_session":
                sessions = auteval["sessions"]
                header_row = ["Critère"] + [
                    f"S{i+1}\n{_format_date(sess.get('date', ''))}"
                    for i, sess in enumerate(sessions)
                ]
                table_data = [header_row]
                for criterion in AUTEVAL_CRITERIA:
                    row = [criterion]
                    for sess in sessions:
                        scores = sess.get("scores", {})
                        row.append("—" if not scores else str(scores.get(criterion, 0)))
                    table_data.append(row)
                n_cols = len(header_row)
                crit_col = 5 * cm
                score_col = (13.5 * cm - crit_col) / max(n_cols - 1, 1)
                col_widths = [crit_col] + [score_col] * (n_cols - 1)
                tbl = Table(table_data, colWidths=col_widths)
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), indigo),
                    ("TEXTCOLOR", (0, 0), (-1, 0), white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [white, HexColor("#F9FAFB")]),
                    ("GRID", (0, 0), (-1, -1), 0.5, lightgrey),
                    ("ALIGN", (1, 0), (-1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]))
                story.append(tbl)
                legend_text = (
                    "<i>Scores de 0 à 5 — 0 : non renseigné · 1 : Faible · 2 : Passable · "
                    "3 : Bien · 4 : Très bien · 5 : Excellent. "
                    "La première séance ne comporte pas d'autoévaluation (—).</i>"
                )
                story.append(Paragraph(legend_text, small_style))
            else:
                scores = auteval["scores"]
                table_data = [["Critère", "Note /5", "Appréciation"]]
                for criterion in AUTEVAL_CRITERIA:
                    val = scores.get(criterion, 0)
                    label = DOT_LABELS[val] if 0 <= val < len(DOT_LABELS) else ""
                    table_data.append([criterion, str(val), label])
                tbl = Table(table_data, colWidths=[6 * cm, 2.5 * cm, 5 * cm])
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), indigo),
                    ("TEXTCOLOR", (0, 0), (-1, 0), white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [white, HexColor("#F9FAFB")]),
                    ("GRID", (0, 0), (-1, -1), 0.5, lightgrey),
                    ("ALIGN", (1, 0), (1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]))
                story.append(tbl)
        elif content:
            for line in content.split("\n"):
                stripped = line.strip()
                if stripped:
                    story.append(Paragraph(stripped, body_style))
        else:
            story.append(Paragraph("<i>[Section non renseignée]</i>", small_style))
        story.append(Spacer(1, 0.3 * cm))

    # Formule de clôture — lieu du patient (tag) ou placeholder rouge si absent
    if patient_label:
        lieu_part = patient_label
    else:
        lieu_part = '<font color="#DC2626">[Lieu]</font>'
    closing_text = f"Merci de votre confiance, fait à {lieu_part}, le {date_str}."
    if therapist_name:
        closing_text += f"<br/>{therapist_name}"
    story.append(Paragraph(closing_text, closing_style))

    doc.build(story, onFirstPage=on_first_page, onLaterPages=on_later_pages)
    return dest


def export(sections: list[dict], template_name: str, patient_name: str = "", patient_id: str = "") -> None:
    session_id = uuid.uuid4().hex[:8]

    # Dossier Bilan/ dans le dossier du patient si possible, sinon dossier général
    patient_dir = patient_dir_for(patient_id) if patient_id else None
    if patient_dir:
        out_dir = patient_dir / "Bilan"
        out_dir.mkdir(parents=True, exist_ok=True)
    else:
        out_dir = _output_dir()

    # Récupérer le tag lieu du patient
    patient_label = ""
    if patient_dir and (patient_dir / "patient.json").exists():
        try:
            data = json.loads((patient_dir / "patient.json").read_text(encoding="utf-8"))
            patient_label = data.get("label", "").strip()
        except Exception:
            pass

    filename = _make_filename(session_id)
    settings = get_settings()

    docx_path: str | None = None
    pdf_path: str | None = None

    if _DOCX_AVAILABLE:
        _emit({"type": "progress", "status": "docx", "message": "Génération du fichier Word…"})
        try:
            docx_path = str(_export_docx(sections, template_name, patient_name, out_dir, filename, settings, patient_label))
        except Exception as e:
            _emit({"type": "error", "message": f"Erreur Word : {e}"})
            return
    else:
        _emit({"type": "warning", "message": "python-docx non disponible."})

    if _PDF_AVAILABLE:
        _emit({"type": "progress", "status": "pdf", "message": "Génération du PDF…"})
        try:
            pdf_path = str(_export_pdf(sections, template_name, patient_name, out_dir, filename, settings, patient_label))
        except Exception as e:
            _emit({"type": "error", "message": f"Erreur PDF : {e}"})
            return
    else:
        _emit({"type": "warning", "message": "reportlab non disponible."})

    # Enregistrement du bilan dans le dossier patient
    if patient_id and (docx_path or pdf_path):
        try:
            record_bilan(patient_id, docx_path or "", pdf_path or "")
        except Exception:
            pass

    _emit({
        "type": "complete",
        "docx_path": docx_path,
        "pdf_path": pdf_path,
        "folder_path": str(out_dir),
        "filename": filename,
    })
